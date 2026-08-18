import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42) # #0f172a
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(5, 150, 105) # #059669
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(37, 99, 235) # #2563eb
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
    return p

def add_body_paragraph(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(11)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(15, 23, 42)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 65, 85) # #334155
    return p

def add_bullet_item(doc, bold_title, description):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_bold = p.add_run(bold_title + ": ")
    r_bold.font.name = 'Calibri'
    r_bold.font.size = Pt(11)
    r_bold.bold = True
    r_bold.font.color.rgb = RGBColor(15, 23, 42)
    
    r_desc = p.add_run(description)
    r_desc.font.name = 'Calibri'
    r_desc.font.size = Pt(11)
    r_desc.font.color.rgb = RGBColor(51, 65, 85)
    return p

def main():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    # --- TITLE COVER HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    
    r_title = title_p.add_run("VotePulse: Secure & Real-Time Voting Management System")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(24)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    r_sub = sub_p.add_run("Minor Project Synopsis & Technical Presentation Document")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(5, 150, 105)
    r_sub.bold = True
    
    # Metadata Table Box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in meta_table.rows:
        for cell in row.cells:
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
    meta_table.cell(0, 0).paragraphs[0].add_run("Project Title: VotePulse Online Voting System").bold = True
    meta_table.cell(0, 1).paragraphs[0].add_run("Target Platform: Web & Progressive Web App (PWA)").bold = True
    meta_table.cell(1, 0).paragraphs[0].add_run("Prepared For: Internal Guide Review").bold = True
    meta_table.cell(1, 1).paragraphs[0].add_run("Core Stack: Node.js, Express, HTML5, CSS3, JS").bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- 1. PROBLEM IDENTIFICATION ---
    add_heading_styled(doc, "1. Problem Identification", 1)
    add_body_paragraph(doc, "Traditional paper-based voting systems and outdated election administration workflows suffer from significant operational, security, and accessibility challenges in modern institutional environments.")
    
    add_bullet_item(doc, "High Operational & Paper Cost", "Manual ballot paper printing, ballot box physical distribution, physical security personnel deployment, and manual vote counting incur substantial financial and temporal expenditure.")
    add_bullet_item(doc, "Vulnerability to Fraud & Double Voting", "Physical ballots are susceptible to identity impersonation, ballot box stuffing, illegal multi-voting, and missing or damaged ballot papers during transit.")
    add_bullet_item(doc, "Delayed Tallying & Human Error", "Manual ballot counting is slow, labor-intensive, and prone to human miscalculations or disputes during election result compilation.")
    add_bullet_item(doc, "Queue Congestion & Limited Accessibility", "In-person voting forces long queues and restricts participation for remote, non-resident, or mobility-impaired voters.")
    add_bullet_item(doc, "Lack of Instant Auditability", "Voters have zero digital verification mechanism to ensure their individual ballot was accurately registered and sealed without tampering.")
    
    # --- 2. PROPOSED APPROACH / SOLUTION ---
    add_heading_styled(doc, "2. Proposed Approach & Solution", 1)
    add_body_paragraph(doc, "VotePulse is an enterprise-grade, lightweight, and secure digital voting management application engineered to streamline elections with maximum transparency and zero friction.")
    
    add_bullet_item(doc, "Real-Time OTP Verification Engine", "Integrates a 6-digit auto-advancing OTP verification system with dynamic SMS/push alert toasts, audio chimes, and a 5-minute security countdown to guarantee authentic voter identification.")
    add_bullet_item(doc, "Strict Server-Side Single-Vote Enforcement", "Implements dual composite validation (election_id + voter_id) at both client and REST API levels, completely blocking any duplicate vote attempts.")
    add_bullet_item(doc, "Sealed Digital Ballot Receipts", "Generates an instant, tamper-proof digital vote receipt featuring a unique cryptographic tracking hash, voter timestamp, and candidate confirmation badge for full post-election verification.")
    add_bullet_item(doc, "Responsive Light & Dark Theme PWA", "Delivers an accessible Progressive Web App (PWA v6.0) interface with light/dark theme switching, dynamic font scaling, and offline Service Worker caching for seamless mobile & desktop usage.")
    add_bullet_item(doc, "Comprehensive Admin Console & Live Analytical Reporting", "Provides administrators with real-time voter metrics, candidate lifecycle controls, one-click poll status toggles, and automated CSV report export functionality.")

    # --- 3. SYSTEM DESIGN & MODULES ---
    add_heading_styled(doc, "3. System Design & Modules", 1)
    add_body_paragraph(doc, "The platform adopts a decoupled, modular Client-Server Architecture consisting of four core operational modules:")
    
    add_heading_styled(doc, "A. Voter Portal Module", 2)
    add_bullet_item(doc, "Authentication & Registration", "Allows voters to log in using their credentials and complete real-time OTP verification.")
    add_bullet_item(doc, "Ballot Exploration & Manifestos", "Displays active election polls with candidate portfolios, party affiliations, and detailed manifestos.")
    add_bullet_item(doc, "Instant Vote Casting", "Renders modal confirmation prompts and securely transmits single-vote requests.")
    add_bullet_item(doc, "Digital Ballot Receipt Display", "Shows an immutable receipt with candidate name, voter ID, timestamp, and receipt hash.")

    add_heading_styled(doc, "B. Administrative Console Module", 2)
    add_bullet_item(doc, "Overview Analytics Dashboard", "Displays live metrics (Total Registered Voters, Active Elections, Candidate Count, Total Votes Cast).")
    add_bullet_item(doc, "Election Management", "Supports creating, activating, pausing, and closing election polls.")
    add_bullet_item(doc, "Candidate Directory", "Enables candidate registration with party affiliations, department labels, and manifesto details.")
    add_bullet_item(doc, "Live Tally & Report Export", "Renders real-time percentage progress bars and exports CSV election reports.")

    add_heading_styled(doc, "C. REST API & Data Persistence Engine", 2)
    add_bullet_item(doc, "Express.js REST Endpoints", "Serves JSON routes (/api/auth, /api/otp, /api/elections, /api/candidates, /api/vote, /api/admin/stats).")
    add_bullet_item(doc, "Lightweight Atomic JSON Storage", "Ensures transactional data integrity and zero data loss using atomic database persistence.")

    add_heading_styled(doc, "D. PWA & Service Worker Offline Layer", 2)
    add_bullet_item(doc, "Offline Service Worker (v6.0)", "Caches core assets (HTML, CSS, JS, Manifest) for uninterrupted offline availability.")
    add_bullet_item(doc, "Cross-Platform Mobile PWA", "Includes manifest.json with install prompts for Android, iOS, and Desktop devices.")

    # --- 4. FEASIBILITY ANALYSIS ---
    add_heading_styled(doc, "4. Feasibility Analysis", 1)
    
    add_bullet_item(doc, "Technical Feasibility", "Built on proven web technologies (Node.js, Express, HTML5, Vanilla CSS3, JS ES6+). High hardware compatibility across all modern web browsers without heavy third-party runtime requirements.")
    add_bullet_item(doc, "Operational Feasibility", "Extremely low learning curve. Intuitive step-by-step UI allows non-technical voters to cast ballots in under 30 seconds. Automated admin workflows eliminate manual ballot handling.")
    add_bullet_item(doc, "Economic Feasibility", "Leverages an entirely open-source technology stack with zero licensing overhead. Can be hosted on zero-cost or minimal-cost web hosting platforms (e.g., GitHub Pages, Vercel, Render).")
    add_bullet_item(doc, "Legal & Security Compliance Feasibility", "Adheres to voter privacy principles by protecting voter identities while maintaining verifiable ballot integrity.")

    # --- 5. ESTIMATED BUDGET ---
    add_heading_styled(doc, "5. Estimated Budget", 1)
    add_body_paragraph(doc, "Due to the selection of open-source technologies and efficient architecture, the overall financial expenditure for development and initial deployment is minimal:")

    budget_table = doc.add_table(rows=5, cols=3)
    budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Component / Service", "Description / Tooling", "Estimated Cost (INR)"]
    for i, h in enumerate(headers):
        cell = budget_table.cell(0, i)
        set_cell_background(cell, "0F172A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    data = [
        ["Development Stack", "Node.js, Express, VS Code, Git, GitHub (Open Source)", "₹0.00"],
        ["Hosting & Domain", "GitHub Pages / Cloud Server (Free Tier)", "₹0.00"],
        ["Security & SSL", "Let's Encrypt / GitHub HTTPS Certificate", "₹0.00"],
        ["Database Storage", "Lightweight JSON File Persistence System", "₹0.00"]
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = budget_table.cell(row_idx, col_idx)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            else:
                set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.add_run(text)
            if col_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 6. DEVELOPMENT PLAN & TIMELINE ---
    add_heading_styled(doc, "6. Development Plan & Timeline", 1)
    add_body_paragraph(doc, "The project was executed following an Agile methodology spanning a 6-week development lifecycle:")

    timeline_table = doc.add_table(rows=7, cols=3)
    timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_headers = ["Phase & Period", "Key Tasks & Deliverables", "Status"]
    for i, h in enumerate(t_headers):
        cell = timeline_table.cell(0, i)
        set_cell_background(cell, "059669")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    t_data = [
        ["Week 1: Requirements & Security", "Problem identification, security spec definition, API design.", "Completed"],
        ["Week 2: Backend REST Engine", "Node.js/Express REST server setup, JSON database architecture.", "Completed"],
        ["Week 3: Real-Time OTP & Voting Engine", "6-digit OTP verification, push alert toast, strict single-vote block.", "Completed"],
        ["Week 4: Frontend UI/UX Design System", "Voter portal, Admin dashboard, theme toggle engine, responsive design.", "Completed"],
        ["Week 5: PWA & Offline Support", "Service Worker caching v6.0, manifest.json, install prompt logic.", "Completed"],
        ["Week 6: Testing & Deployment", "UAT testing, bug fixes, GitHub repository & Pages deployment.", "Completed"]
    ]
    
    for row_idx, row_data in enumerate(t_data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = timeline_table.cell(row_idx, col_idx)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            else:
                set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            if col_idx == 2:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 7. EXPECTED OUTCOME ---
    add_heading_styled(doc, "7. Expected Outcome", 1)
    add_body_paragraph(doc, "The implementation of the VotePulse system yields the following direct outcomes and benefits:")

    add_bullet_item(doc, "Zero Duplicate Votes", "Guaranteed single-vote integrity through strict server-side validation.")
    add_bullet_item(doc, "Instant Election Results", "Elimination of manual counting delays; results tallied automatically in real time.")
    add_bullet_item(doc, "95% Reduction in Administrative Expenses", "Complete paperless voting removes printing, logistics, and manual processing costs.")
    add_bullet_item(doc, "Verifiable Voter Transparency", "Every voter receives a digital receipt with confirmation of their candidate selection.")
    add_bullet_item(doc, "Universal Accessibility", "Seamless accessibility across mobile devices, tablets, and desktop computers via PWA v6.0.")

    # Save document
    out_path = os.path.join("d:\\Minor project", "VotePulse_Project_Report.docx")
    doc.save(out_path)
    print(f"Report document successfully created at: {out_path}")

if __name__ == "__main__":
    main()
